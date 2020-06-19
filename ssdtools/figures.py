import gc
import os
import io

from pathlib import Path

import numpy as np
import pandas as pd
from scipy.spatial import distance

import matplotlib    ###TODO is dit nodig?
import matplotlib.pyplot as plt
from matplotlib.lines import Line2D
from matplotlib.colors import ListedColormap, to_rgba, Normalize
from matplotlib.patches import Patch, Rectangle
from matplotlib import ticker
from matplotlib import colorbar

from imageio import imread
from descartes import PolygonPatch
from geopandas import GeoDataFrame

from ssdtools import branding ###TODO Is dit nodig
from ssdtools.branding import default
from ssdtools.traffic import Traffic
from ssdtools.traffic import TrafficAggregate
from ssdtools.traffic import read_file
from ssdtools.grid import Grid, read_grid

# For Docx figures and tables
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

dir = os.path.dirname(__file__)

def fig_to_word(fig,
                table,
                width=Inches(16.48/2.54),
                dpi=600):
    
   
    if isinstance(fig, str):
        z = fig
    else:
        # Create an in-memory stream for text I/O to capture the image
        z = io.BytesIO()
        
        # Save the figure to the in-memory stream
        fig.savefig(z, 
                    dpi=dpi,
                    bbox_inches='tight',
                    pad_inches=0)
        
    # Select the correct paragraph of the document
    p = table.rows[0].cells[0].paragraphs[0]
    
    # Add the picture to the paragraph and align the image
    p.add_run().add_picture(z,
                            width=width)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if not isinstance(fig, str):
        z.close()


def fill_table(docx_table, dataframe, column_header=True, row_header=True, column_header_style=None,
               row_header_style=None, body_style=None):
    column_start = 0
    row_start = 0

    if column_header and row_header:

        # Set the row and column start
        row_start = 1
        column_start = 1

        # Set the first cell
        cell = docx_table.cell(0, 0)
        cell.paragraphs[0].add_run('{}'.format(dataframe.index.name))

        if column_header_style is not None:
            cell.paragraphs[0].style = column_header_style

    if column_header:

        # Set the row start
        row_start = 1

        # Set the column names
        for cell, header in zip(docx_table.rows[0].cells[column_start:], dataframe.columns.values):
            cell.paragraphs[0].add_run(str(header))

            if column_header_style is not None:
                cell.paragraphs[0].style = column_header_style

    if row_header:

        # Set the column start
        column_start = 1

        # Set the row names
        for cell, header in zip(docx_table.columns[0].cells[row_start:], dataframe.index.values):
            cell.paragraphs[0].add_run(str(header))

            if row_header_style is not None:
                cell.paragraphs[0].style = row_header_style

    # Set the tables content
    for j in range(dataframe.shape[0]):
        for i in range(dataframe.shape[1]):
            cell = docx_table.cell(j + row_start, i + column_start)
            value = dataframe.iloc[j, i]
            if isinstance(value, str) and value != 'nan':
                cell.paragraphs[0].add_run(value)
            elif isinstance(value, float) and not np.isnan(value):
                ###TODO is dit centraal handig
                cell.paragraphs[0].add_run('{0:,.0f}'.format(value).replace(',', '.'))

            if body_style is not None:
                cell.paragraphs[0].style = body_style


def Formatter_1000sep0d(x, pos):
    'The two args are the value and tick position'
    return '{:,.0f}'.format(x).replace(',', '.')

def soften_colormap_edge(colormap, transition_width=.25, alpha=1.):
    """
    Soften the colormap by applying a linear transition to zero at the front of the colormap.

    # # # # # # # # # < 1
    #   ----------- # < alpha
    #  /            #
    # /             #
    # # # # # # # # # < 0
    ^   ^           ^
    0   |           1
        |
    transition_width

    :param ColorMap colormap: the colormap to soften.
    :param float transition_width: the width (as percentage) of the transition, should be between 0 and 1.
    :param float alpha: the maximum alpha. Should have a value between 0 and 1.
    :rtype: ColorMap
    """
    # Get the colormap colors
    colormap_colors = colormap(np.arange(colormap.N))

    # Get the transition range
    transition_range = np.arange(int(colormap.N * transition_width))

    # Set alpha
    colormap_colors[:, -1] = alpha

    # Set the alpha in the transition range
    colormap_colors[transition_range, -1] *= np.linspace(0, 1, transition_range.shape[0])

    # Create new colormap
    return ListedColormap(colormap_colors)


def soften_colormap_center(colormap, alpha=1.):
    """
    Soften the colormap by applying a linear transition from 1 to 0 at the first half and from 0 to 1 at the second
    half.

    # # # # # # < 1
    # \     / # < alpha
    #  \   /  #
    #   \ /   #
    # # # # # # < 0
    ^    ^    ^
    0   0.5   1

    :param ColorMap colormap: the colormap to soften.
    :param float alpha: the maximum alpha. Should have a value between 0 and 1.
    :rtype: ColorMap
    """
    # Get the colormap colors
    colormap_colors = colormap(np.arange(colormap.N))

    # Calculate the length of half the list
    n_2 = int(colormap.N / 2)

    # Set alpha
    colormap_colors[:, -1] = alpha

    # Set the alpha in the transition range
    colormap_colors[:n_2, -1] *= np.linspace(1, 0, n_2)
    colormap_colors[n_2:, -1] *= np.linspace(0, 1, colormap.N - n_2)

    # Create new colormap
    return ListedColormap(colormap_colors)


class GridPlot(object):
    def __init__(self,
                 grid,
                 other_grid=None,
                 title=None,
                 branding=default,
                 figsize=(21 / 2.54, 21 / 2.54),
                 xlim=(80000, 140000),
                 ylim=(455000, 515000),
                 background=dir + '/data/Schiphol_RD900dpi.png',
                 schiphol_border=dir + '/data/2013-spl-luchtvaartterrein.shp',
                 place_names=dir + '/data/plaatsnamen.csv',
                 scalebar=True,
                 extent=[80000, 158750, 430000, 541375]):

        # Create an ID
        self.id = str(pd.Timestamp.utcnow())

        # Set the grids
        self.grid = grid
        self.other_grid = other_grid

        # Set the plotting options
        self.title = title
        self.branding = default
        self.figsize = figsize
        self.xlim = xlim
        self.ylim = ylim
        self.background = background
        self.schiphol_border = schiphol_border
        self.place_names = place_names
        self.extent = extent

        # Create a new figure
        self.fig, self.ax = self.create_figure()

        # Add background, place names, schiphol border and scalebar
        if background is not None: self.add_background()
        if place_names is not None: self.add_place_names()
        if schiphol_border is not None: self.add_schiphol_border()
        if scalebar: self.add_scalebar()
        
        # Create a placeholder for contour plots
        self.contour_plot = None

    def create_figure(self):
        """
        Initialize a new contour plot.

        :return: the figure and axis handles.
        """

        # Create a new figure with figsize
        fig = plt.figure(num=self.id, figsize=self.figsize)

        # Don't show axis
        ax = plt.Axes(fig, [0., 0., 1., 1.])
        ax.set_axis_off()
        fig.add_axes(ax)

        # Square X- and Y-axis
        ax.set_aspect('equal', 'box')

        # Zoom in
        ax.set_xlim(self.xlim)
        ax.set_ylim(self.ylim)

        return fig, ax

    def add_background(self, background=None):
        """
        Add the background map to the figure.

        :param str|np.ndarray background: path to the background file or background image as NumPy array.
        """
        background = self.background if background is None else background
        if isinstance(background, str):
            background = imread(background)
            # self.background = Image.open(self.background)
        self.ax.imshow(background, zorder=0, extent=self.extent)
       
        return self

    def add_place_names(self, place_names=None):
        ###TODO color etc uit plot_style
        """
        Add the place names to the map.
        :param str|pd.DataFrame place_names: path to the place name file or a data frame containing the place names with
        corresponding coordinates.
        :param tuple color: the color of the text.
        """
        place_names = self.place_names if place_names is None else place_names

        if isinstance(place_names, str):
            place_names = pd.read_csv(place_names, comment='#')

        for index, row in place_names.iterrows():
            self.ax.annotate(row['name'],
                             xy=(row['x'], row['y']),
                             **branding.xParams['placenames'])
        return self
        
    def add_schiphol_border(self, schiphol_border=None, background=None):
        ###TODO color etc uit plot_style
        """

        :param str|GeoDataFrame schiphol_border: path to the schiphol_border file or a Pandas DataFrame with a geometry column.
        """
        background = self.background if background is None else background
        if isinstance(background, str):
            background = imread(background)
        
        schiphol_border = self.schiphol_border if schiphol_border is None else schiphol_border
        if isinstance(schiphol_border, str):
            schiphol_border = GeoDataFrame.from_file(schiphol_border)

        fc = 'none'
        ec = (.0, .0, .0)
        poly = schiphol_border['geometry'][0]
        patch = PolygonPatch(poly, fc=fc, ec=ec, lw=0.2, zorder=10)
        self.ax.add_patch(patch)

        # Show background inside the border
        if background is not None:
            im = self.ax.imshow(background, clip_path=patch, clip_on=True, zorder=9, extent=self.extent)
            im.set_clip_path(patch)

        return self

    def add_contourlabels (self,
                           cs,
                           ankerpoint=(110000, 495500),
                           dxy=(7000, 3000)):
    ###TODO beschrijving van de functie
        
        for i, level in enumerate(cs.levels):    
    
            # Shift ankerpoint for level
            d = (58 - level) * 200
            ap = [ankerpoint[0]+d, ankerpoint[1]-d]
            # self.ax.scatter(ap[0], ap[1], color='red', marker='o', s=3)
            
            # Nearest datapunt
            xy2 = [vertices for path in cs.collections[i].get_paths() for vertices in path.vertices]
            p2 = xy2[distance.cdist([ap],xy2).argmin()]
            textloc = p2 + dxy
            
            self.ax.scatter(p2[0],p2[1], 
                            **branding.xParams['contourlabelmarker'],
                            zorder=10)
            
            self.ax.annotate(f'{level:.0f} dB(A)',
                             xy=p2,
                             xytext=textloc,
                             **branding.xParams['contourlabel'])
        return self

    
    def add_scalebar(self, ticks=None):
        """

        :param list(float) ticks: the ticks to use as scale, in km.
        """

        # Scale, with xpos and ypos as position in fig-units
        xpos = .95
        ypos = 0.04
        ticks = [0, 2, 5, 10] if ticks is None else ticks

        # Transform from data to fig
        l_data = ticks[-1] * 1000
        l_disp = self.ax.transData.transform((l_data, 0)) - self.ax.transData.transform((0, 0))
        inv_fig = self.fig.transFigure.inverted()
        l_fig = inv_fig.transform(l_disp)[0]

        ax2 = self.fig.add_axes([xpos - l_fig, ypos, l_fig, 0.01])

        # Transparent background
        ax2.patch.set_alpha(0)

        # Disable spines
        for loc in ['right', 'left', 'top']:
            ax2.spines[loc].set_color('none')

        # geen verticale  gridlines
        ax2.xaxis.grid(which='both', color='None')  

        # Remove Y-as
        ax2.yaxis.set_ticks([])

        # Add X-ticks
        ax2.xaxis.set_ticks(ticks)
        labels = ['{}'.format(tick) for tick in ticks]
        labels[-1] = '{} km'.format(ticks[-1])
        ax2.set_xticklabels(labels)

        # Format
        ax2.spines['bottom'].set_color(get_cycler_color(0))
        ax2.spines['bottom'].set_linewidth(0.5)
        ax2.tick_params(axis='x', labelsize=10, colors=get_cycler_color(0), length=4, direction='in', width=0.5)

        return self


    def add_colorbar(self, contour_plot=None,cax_position=None):

        # Use the contour plot of this object if no contour plot is provided
        contour_plot = self.contour_plot if contour_plot is None else contour_plot

        # Create new axis for the colorbar in the top-right corner. The sequence is left, bottom, width and height.
        cax = self.fig.add_axes([0.9, 0.67, 0.05, 0.3]) if cax_position is None else self.fig.add_axes(cax_position)

        # Add the colorbar
        return (colorbar.ColorbarBase(cax,
                                      cmap=contour_plot.get_cmap(),
                                      norm=Normalize(*contour_plot.get_clim()))
                        .ax.tick_params(labelsize=8,
                                        colors=get_cycler_color(0)))
    

    def add_bandwidth(self,
                      levels=[48,58],
                      mean='mean',
                      labels=['gemiddeld', 'weersinvloeden'],  ###TODO None=geen label
                      colors=None,
                      alpha=0.4,
                      refine_factor=10):
        """
        Add a contour of the grid at the specified noise level. When a multigrid is provided, the bandwidth of the contour
        will be shown.

        :param float levels: the noise level of the contour to plot.
        :param str mean: Average grid type to use for the average noise levels
        :param str labels: List labels for average noise level and bandwith in the legend
        :param Matplotlib-colors colors: List of two colors for the mean and area-contour. If None then first two colors from the default colorcycle are used
        :param float alpha: Transparancy for bandwidth
        :param integer refine_factor: Factor to refine the grid with bi-cubic spline interpolation
        :return:
        """

        # Convert to list
        if not isinstance(levels, list): levels = [levels]
    
        # Default colors
        if colors is None:
            colors = [get_cycler_color(i) for i in range(2)]
        
        # Bandwidth color
        colormap = ListedColormap([colors[1]])

        # Get statistics and refine
        mean = self.grid.statistics()[mean].refine(refine_factor)
        dlo = self.grid.statistics()['dlo'].refine(refine_factor)
        dhi = self.grid.statistics()['dhi'].refine(refine_factor)

        # Extract the x and y coordinates
        x = mean.shape.get_x_coordinates()
        y = mean.shape.get_y_coordinates()
        
        # Select this plot as active figure
        self.select()

        for level in levels:
            # Plot the contour area
            area_mask = np.logical_or(dhi.data < level, dlo.data > level)
            area_grid = np.ma.array(mean.data, mask=area_mask)
            ###TODO meshgrid is niet nodig maar is het soms sneller?
            # plt.contourf(*np.meshgrid(x, y), area_grid, cmap=colormap, alpha=0.4)
            self.ax.contourf(x, y, area_grid, cmap=colormap, alpha=0.4)
   
        # Plot the contours of the statistics
        cs = self.ax.contour(x, y, mean.data, levels=levels, colors=colors[0], linewidths=1)
        self.ax.contour(x, y, dhi.data, levels=levels, colors=colors[1], linewidths=0.5)
        self.ax.contour(x, y, dlo.data, levels=levels, colors=colors[1], linewidths=0.5)
    
        # Add contour labels           
        self.add_contourlabels(cs=cs)
        
        # Add legend
        cs1 = Line2D([], [], color=colors[0], marker='None')
        cs2 = Patch(fc=to_rgba(colors[1], alpha=0.4), ec=colors[1], lw=0.5)
        self.ax.legend(handles=[cs1, cs2],
                       labels=labels,
                       title=r'Geluidbelasting $L_{' + self.grid.unit[1:] + r'}$',
                       **branding.xParams['contourlegend'])                        
        return 



    def add_contours(self,
                     levels,
                     other_grid=False,
                     colors=None,
                     contourlabels=True,
                     refine_factor=10):
        """
        Add a contour of the grid at the specified noise level. When a multigrid is provided, the bandwidth of the contour
        will be shown.

        :param float levels: the noise levels for the contours to plot.
        :param colors: Contourline color.
        :param boolean labels: Add a label with the noise level to the contour 
        :param integer refine_factor: Factor to refine the grid with bi-cubic spline interpolation    
        :return:
        """

        # Convert to list
        if not isinstance(levels, list): levels = [levels]

        # Default color (same for all levels)
        if colors is None:
            colors = get_cycler_color(int(other_grid))
            
        # Select this plot as active figure
        self.select()

        # Refine the grid
        if not other_grid:
            grid = self.grid.refine(refine_factor)
        else:
            grid = self.other_grid.refine(refine_factor)        

        # Extract the x and y coordinates
        x = grid.shape.get_x_coordinates()
        y = grid.shape.get_y_coordinates()

        cs = self.ax.contour(x, 
                             y, 
                             grid.data, 
                             levels=levels, 
                             colors=colors, 
                             linewidths=1)
          
        # Add contour labels           
        if contourlabels:
            self.add_contourlabels(cs=cs)

        return 


    ###TODO wat is het verschil met add_contours???
    def add_individual_contours(self, level, primary_color=None, secondary_color=None, refine_factor=20):
        """
        Add a contour of the grid at the specified noise level. When a multigrid is provided, all contours of the
        individual grids will be shown.

        :param float level: the noise level of the contour to plot.
        :param primary_color: color for the main contour.
        :param secondary_color: color for the secondary contours (only used for multigrids).
        :return:
        """

        # Select this plot as active figure
        self.select()

        # Refine the grid
        grid = self.grid.copy().refine(refine_factor)

        # Extract the x and y coordinates
        x = grid.shape.get_x_coordinates()
        y = grid.shape.get_y_coordinates()

        # If the grid is a multigrid, all noise levels should be plotted.
        if isinstance(grid.data, list):

            # Get the various statistics of the data
            statistic_grids = self.grid.statistics()

            # Extract the data of interest
            mean_grid = statistic_grids['mean'].resize(grid.shape)
            dhi_grid = statistic_grids['dhi'].resize(grid.shape)
            dlo_grid = statistic_grids['dlo'].resize(grid.shape)

            # Plot all individual contours
            for year_data in grid.data:
                self.ax.contour(x,
                                y,
                                year_data,
                                levels=[level],
                                colors=secondary_color,
                                linewidths=3,
                                alpha=0.1)

            # Plot the contours of the statistics
            ###Ed-test
            cs = self.ax.contour(x, y, mean_grid.data, levels=[level], colors=primary_color, linewidths=1) #[1, 1])
            self.ax.contour(x, y, dhi_grid.data, levels=[level], colors=secondary_color, linewidths=0.5) #[0.5, 0.5])
            self.ax.contour(x, y, dlo_grid.data, levels=[level], colors=secondary_color, linewidths=0.5) #[0.5, 0.5])

        # The input is a single grid, so only a single contour should be plotted
        else:
            cs = self.ax.contour(x, y, grid.data, levels=[level], colors=primary_color, linewidths=1) #[1, 1])

        return cs


    def add_heatmap(self,
                    colormap=None,
                    soften_colormap=True,
                    alpha=0.4,
                    refine_factor=10,
                    **kwargs):
        """
        Show a grid by creating a heatmap.

        :param ColorMap colormap: the colormap to apply. Default is 'viridis'
        :param bool soften_colormap: soften the colormap by making the edge transparent.
        :param float alpha: the maximum alpha. Should have a value between 0 and 1.
        :param int refine: a multiplication factor for the number of additional layers to plot, most colormaps consist
        of 64 colors.
        :param kwargs: optional arguments for the underlying contourf function.
        :return:
        """

        # Default colors
        ###TODO werkt dit?
        if colormap is None:
            colormap = plt.rcParams['image.cmap']
            
        # Select this plot as active figure
        self.select()

        # Refine the grid
        grid = self.grid.copy().refine(refine_factor)

        # Extract the x and y coordinates
        x = grid.shape.get_x_coordinates()
        y = grid.shape.get_y_coordinates()

        # Add the transparency to the colormap
        if soften_colormap:
            colormap = soften_colormap_edge(colormap, transition_width=.25, alpha=alpha)

        # Plot the contour area
        self.contour_plot = self.ax.contourf(*np.meshgrid(x, y), grid.data, levels=refine_factor * colormap.N, cmap=colormap,
                                             **kwargs)

        return self.contour_plot


    def add_comparison_heatmap(self, 
                               other_grid,
                               deltas=[-3,3],
                               colormap=None,
                               soften_colormap=True,
                               alpha=1.0,
                               method='energetic',
                               positive_scale=False,
                               refine_factor=10,
                               **kwargs):
        """
        Compare two grids by creating a heatmap.

        :param Grid other_grid: the noise grid to compare.
        :param ColorMap colormap: the colormap to apply. If None the default from branding is used
        :param bool soften_colormap: soften the colormap by making the center transparent.
        :param float alpha: the maximum alpha. Should have a value between 0 and 1.
        :param kwargs: optional arguments for the underlying contourf function.
        :return:
        """

        # Select this plot as active figure
        self.select()

        # Default colors
        if colormap is None:
            colormap = branding.xParams['cmap_diff']
            
        # Align the shape of the other grid to the original grid
        diff_grid = other_grid.copy().resize(self.grid.shape)

        #compute scaling (energetically scale)
        if self.grid.unit == 'Lden':
            threshold = 48
        elif self.grid.unit == 'Lnight': 
            threshold = 40      
        elif self.grid.unit == 'LAmax': 
            threshold = 75  
        elif self.grid.unit == 'NAxx':
            threshold =0
            
        scale                               = np.ones(diff_grid.data.shape)
        scale[diff_grid.data<threshold]     = 10**((diff_grid.data[diff_grid.data<threshold]-threshold)/10)
        
        # Subtract the original grid from the other grid
        diff_grid.data -= self.grid.data   
        
        if method =='energetic':
            diff_grid.data *= scale

        # Refine the grid
        diff_grid.refine(refine_factor)

        # Extract the x and y coordinates
        x = diff_grid.shape.get_x_coordinates()
        y = diff_grid.shape.get_y_coordinates()

        # Add the transparency to the colormap
        if soften_colormap:
            colormap = soften_colormap_center(colormap,
                                              alpha=alpha)
            
        if soften_colormap and positive_scale:
            colormap = soften_colormap_edge(colormap,
                                            transition_width=1,
                                            alpha=alpha)

        # Plot the contour area
        self.contour_plot = self.ax.contourf(*np.meshgrid(x, y),
                                             diff_grid.data,
                                             levels=colormap.N,
                                             cmap=colormap,
                                             vmin=deltas[0],
                                             vmax=deltas[1],
                                             **kwargs)

        return self.contour_plot
    
    def select(self):
        plt.figure(self.id)
        plt.sca(self.ax)

    def save(self, *args, **kwargs):
        return self.fig.savefig(*args, **kwargs)

    def show(self):
        return self.fig.show()


def table_denem_traffic(traffic,
                        traffic_kwargs={'sep':None, 'engine':'python'},
                        index={'D':'dag',
                               'E':'avond',
                               'N':'nacht',
                               'EM':'vroege ochtend'},
                        columns={'L':'landingen',
                                 'T':'starts'},
                        total='totaal',
                        period='periode',
                        wordtable=None,
                        style={}):
    """
    Create a table with traffic per denem-period.
    
    :param str|pd.DataFrame|TrafficAggregate traffic: traffic data containing VVC-code in d_ac_cat or c_ac_cat and movements in total 
    :param dict taffic_kwargs: kwargs for reading traffic data
    :param dict index: dict for renaming the index used for the ticklabels
    :param dict columns: dict for renaming the columns used for legend labels
    :return: return a pd.DataFrame.
    """
    
    # Read traffic
    if isinstance(traffic, str):
        traffic = Traffic.read_daisy_phase_file(traffic, **traffic_kwargs)
    
    # Get the day (D), evening (E), night (N) distribution
    df = traffic.get_denem_distribution(separate_by='d_lt')
       
    # Add totals
    df[total] = df.sum(axis=1) 
    df = df.append(pd.DataFrame({total: df[total].sum()}, index=[total]))
       
    # Round to hundreds
    df = df.round(-2)
    
    # Rename the columns and index
    df = df.rename(index=index, columns=columns)
    df.index.name = period

    # Print table
    with pd.option_context('display.float_format', '{:.0f}'.format):    
        print(df.fillna(''))
        
    if wordtable:
        ###TODO Opmaak
        fill_table(wordtable, df, **style)

    return df


def table_season_traffic(traffic,
                         traffic_kwargs={'sep':None, 'engine':'python'},
                         labels=['winter', 'zomer'],
                         index={'L':'landingen',
                                'T':'starts'},
                         columns={'D':'dag',
                                  'E':'avond',
                                  'N':'nacht',
                                  'EM':'vroege ochtend'}):
    """
    Create a table with traffic per season. The table can be used as input for plot_season_traffic
    
    :param str|pd.DataFrame|TrafficAggregate traffic: traffic data containing VVC-code in d_ac_cat or c_ac_cat and movements in total 
    :param dict taffic_kwargs: kwargs for reading traffic data
    :param str labels: List with strings to identify the traffics 
    :param dict index: dict for renaming the index used for the ticklabels
    :param dict columns: dict for renaming the columns used for legend labels
    :return: return a pd.DataFrame.
    """
 
    # converteer naar list
    if not isinstance(traffic, list): traffic = [traffic]
    if not isinstance(labels, list): labels = [labels]
    
    if len(traffic) != len(labels):
        print('Error: Traffic and labels do not have equal dimensions')
            
    # List to store results
    results = []
    
    for trf, label in zip(traffic, labels):
        # Read traffic
        if isinstance(trf, str):
            trf = Traffic.read_daisy_phase_file(trf, **traffic_kwargs)
        
        # Get the day (D), evening (E), night (N) distribution
        results.append(trf.get_denem_distribution(separate_by='d_lt').round(-2).T)
            
    # Store results
    df = pd.concat(results, keys=labels)

        
    # Rename the columns and index
    df = df.rename(index=index, columns=columns)

    # Print table
    with pd.option_context('display.float_format', '{:.0f}'.format):    
        print(df)

    return df
    
    
def plot_season_traffic(table=None,
                        traffic=None,
                        traffic_kwargs={'sep':None, 'engine':'python'},
                        labels=None,
                        xlabel='vliegtuigbewegingen',
                        xlim=(0,160000),
                        tickformat=Formatter_1000sep0d,
                        ncol=None,
                        dpi=600,
                        fname='',
                        figsize=(8.27, 2.76),
                        wordtable=None,
                        **kwargs):
    """
    A function to create a traffic per season plot. Can also be used to plot other grouped horizontal stacked bar
    charts.
    
    :param pd.DataFrame table: a dataframe containing the numbers to visualise. The dataframe should have a
    2-level multiindex, where the first level is the seasons and the second level is the type of operation. The columns
    are used as labels for the data.
    :param str|pd.DataFrame|TrafficAggregate traffic: traffic data containing VVC-code in d_ac_cat or c_ac_cat and movements in total 
    :param dict taffic_kwargs: kwargs for reading traffic data
    :param str labels: List with strings to identify the traffics 
    :param str xlabel: label for the x-axis
    :param set xlim: range for the x-axis
    :param str|function tickformat: format string or function for ticklabels    
    :param int ncol: number of columns in legend
    :param int dpi: dpi for saving figure to file, default is 600
    :param str fname: Name for the file, default is '' and no fig will be saved
    :param set figsize: Figsize in inches, default (21/2.54, 7/2.54)
    :return: if fname='', return a Matplotlib figure and axes.
    """
    
    if table is None:
        table = table_season_traffic(traffic=traffic,
                                     traffic_kwargs=traffic_kwargs,
                                     labels=labels)   
    # plot
    return plot_barh(table,
                     stacked=True,
                     xlabel=xlabel,
                     xlim=xlim,
                     tickformat=tickformat,
                     figsize=figsize,
                     dpi=dpi,
                     fname=fname,
                     wordtable=wordtable,
                     **kwargs)
            

def plot_bar(table=None,
             xlabel=None,
             ylabel=None,
             tickformat=None,
             ncol=None,
             fname='',
             figsize=(8.27, 2.76),
             dpi=600,
             wordtable=None,
             **kwargs):
    """
    A function to create a barplot. 

    :param pd.DataFrame table: Dataframe to plot, see table_aircraft_types   
    :param str xlabel: label for the x-axis
    :param str ylabel: label for the y-axis
    :param str|function tickformat: format string or function for ticklabels
    :param int ncol: number of columns in legend
    :param str fname: Name for the file, default is '' and no fig will be saved
    :param set figsize: Figsize in inches, default (21/2.54, 7/2.54)
    :param int dpi: dpi for saving figure to file, default is 600
    :return: if fname='', return a Matplotlib figure and axes.
    """
        
    # plot format
    plot_format = branding.xParams['barplot']     # general settings
    plot_format.update(edgecolor='none', rot=0)   # plot specific
    plot_format.update(kwargs)                    # user input                                   

    # plot
    ax = table.plot.bar(figsize=figsize,
                        **plot_format)

    # geen verticale  gridlines
    ax.xaxis.grid(which='major', color='None')   

    # X-as
    if len(table.index.names) == 2:
        ax.set_xticklabels(table.index.get_level_values(1))
        branding.set_xlabels(table.index.get_level_values(0).unique(), ax=ax)
    elif xlabel is not None:
        branding.set_xlabels(xlabel, ax=ax)
    else:
        ax.set_xlabel('') # verberg as-label
        
    # Y-as
    if tickformat is not None:
        if isinstance(tickformat, str):
            ax.yaxis.set_major_formatter(ticker.StrMethodFormatter(tickformat))
        elif callable(tickformat):
            ax.yaxis.set_major_formatter(ticker.FuncFormatter(tickformat))
        else:
            print('Ignore unknown tickformat')
    
    if ylabel is not None:
        branding.set_ylabels(ylabel, ax=ax)

    # legend
    if ncol is None: ncol = len(table.columns)
    leg = ax.legend(ncol=ncol,
                    handletextpad=-0.5,
                    **branding.xParams['legend'])

    for patch in leg.get_patches():  # Maak de patches vierkant
        patch.set_height(5)
        patch.set_width(5)
        patch.set_y(-1)              # Vertikaal uitlijnen
        
    # save figure
    fig = plt.gcf()
    if fname:
        fig.savefig(fname, dpi=dpi)
    
    # Export figure to Word
    if wordtable:
        fig_to_word(fig=fig,
                    table=wordtable,
                    dpi=dpi)
        
    if not (fname or wordtable):
        return fig, ax 
    else:
        plt.close(fig)
        return


def plot_barh(table=None,
              xlabel=None,
              ylabel=None,
              tickformat=None,
              ncol=None,
              fname='',
              figsize=(8.27, 2.76),
              dpi=600,
              wordtable=None,
              **kwargs):
    """
    A function to create a barh-plot. 

    :param pd.DataFrame table: Dataframe to plot, see table_aircraft_types   
    :param str xlabel: label for the x-axis
    :param str ylabel: label for the y-axis
    :param str|function tickformat: format string or function for ticklabels
    :param int ncol: number of columns in legend
    :param str fname: Name for the file, default is '' and no fig will be saved
    :param set figsize: Figsize in inches, default (21/2.54, 7/2.54)
    :param int dpi: dpi for saving figure to file, default is 600
    :return: if fname='', return a Matplotlib figure and axes.
    """
        
    # plot format
    plot_format = branding.xParams['barplot']     # general settings
    plot_format.update(edgecolor='none', rot=0)   # plot specific
    plot_format.update(kwargs)                    # user input                                   

    # plot
    ax = table.plot.barh(figsize=figsize,
                         **plot_format)

    # geen horizontale  gridlines
    ax.yaxis.grid(which='major', color='None')        

    # X-as
    if tickformat is not None:
        if isinstance(tickformat, str):
            ax.xaxis.set_major_formatter(ticker.FormatStrFormatter(tickformat))
        elif callable(tickformat):
            ax.xaxis.set_major_formatter(ticker.FuncFormatter(tickformat))
        else:
            print('Ignore unknown tickformat')

    if xlabel is not None:
        branding.set_xlabels(xlabel, ax=ax)
    else:
        ax.set_xlabel('') # verberg as-label
        
    # Y-as
    ax.invert_yaxis()
    if len(table.index.names) == 2:
        ax.set_yticklabels(table.index.get_level_values(1),
                           va='center',
                           rotation='vertical')
        branding.set_ylabels(table.index.get_level_values(0).unique()[::-1], ax=ax)
    elif ylabel is not None:
        branding.set_ylabels(ylabel, ax=ax)

    # legend
    if ncol is None: ncol = len(table.columns)
    leg = ax.legend(ncol=ncol,
                    handletextpad=-0.5,
                    **branding.xParams['legend'])

    for patch in leg.get_patches():  # Maak de patches vierkant
        patch.set_height(5)
        patch.set_width(5)
        patch.set_y(-1)              # Vertikaal uitlijnen
        
    # save figure
    fig = plt.gcf()
    if fname:
        fig.savefig(fname, dpi=dpi)

    # Export figure to Word
    if wordtable:
        fig_to_word(fig=fig,
                    table=wordtable,
                    dpi=dpi)

    if not (fname or wordtable):
        return fig, ax 
    else:
        plt.close(fig)
        return


def table_aircraft_types(traffic,
                         traffic_kwargs={'sep':None, 'engine':'python'},
                         labels='handelsverkeer',
                         weight_classes= {0: '< 6',
                                          1: '6 - 40',
                                          2: '6 - 40',
                                          3: '40 - 60',
                                          4: '60 - 160',
                                          5: '60 - 160',
                                          6: '160 - 230',
                                          7: '230 - 300',
                                          8: '> 300',
                                          9: '> 300'}):
    """
    Create a table with the number of aircraft per weight class. The weight classes are based on the VVC-code.
    The table can be used as input for plot_aircraft_types 

    :param str|pd.DataFrame|TrafficAggregate traffic: traffic data containing VVC-code in d_ac_cat or c_ac_cat and movements in total 
    :param dict taffic_kwargs: kwargs for reading traffic data
    :param str labels: List with strings to identify the traffics 
    :param dict weight_classes: Text per VVC weight class to aggregate the traffic
    :return: return a pd.DataFrame.
    """
         
    # converteer naar list
    if not isinstance(traffic, list): traffic = [traffic]
    if not isinstance(labels, list): labels = [labels]
    
    if len(traffic) != len(labels):
           print('Error: Traffic and labels do not have equal dimensions')
 
    # df for storing results with unique index   
    df = (pd.DataFrame(data=weight_classes.values(), columns=['mtow'])
            .drop_duplicates()
            .set_index('mtow'))
    
    for trf, label in zip(traffic, labels):
        # Read traffic
        ###TODO gaat dit altijd goed?
        if isinstance(trf, TrafficAggregate):
            trf = trf.data
        else:
            trf = read_file(trf, **traffic_kwargs)
    
        # Realisatietraffic aanpassen
        if 'C_VVC' in trf:
            if trf['C_VVC'].isna().any():
                missing = trf['C_VVC'].isna().sum()
                print(f"Info: {missing} missing VVC's in traffic '{label}'")
            trf = (trf.dropna(subset=['C_VVC'])
                      .rename(columns={'C_VVC':'d_ac_cat'})
                      .assign(total=1))
                    
        # mtow class
        mtow = trf['d_ac_cat'].str.get(0).astype(int).map(weight_classes)

        # Group by mtow class
        p = trf.groupby(mtow)['total'].sum()
        p = 100 * p / p.sum()
        
        # Store results
        df = df.join(p, how='left').fillna(0).rename(columns={'total': label})

    # Print table
    with pd.option_context('display.float_format', '{:,.1f}%'.format):    
        print(df.reset_index().to_string(index=False))

    return df


def plot_aircraft_types(table=None,
                        traffic=None,
                        traffic_kwargs={'sep':None, 'engine':'python'},
                        labels='handelsverkeer',
                        weight_classes= {0: '< 6',
                                         1: '6 - 40',
                                         2: '6 - 40',
                                         3: '40 - 60',
                                         4: '60 - 160',
                                         5: '60 - 160',
                                         6: '160 - 230',
                                         7: '230 - 300',
                                         8: '> 300',
                                         9: '> 300'},
                        xlabel='maximum startgewicht in tonnen',
                        ylabel='aandeel in de vloot',
                        tickformat='{x:.0f}%',
                        ylim=(0,60),
                        fname='',
                        figsize=(8.27, 2.76),
                        dpi=600,
                        wordtable=None,
                        **kwargs):
    """
    A function to create a fleetmix plot. 

    :param pd.DataFrame table: Dataframe to plot, see table_aircraft_types   
    :param str|pd.DataFrame|TrafficAggregate traffic: traffic data containing VVC-code in d_ac_cat or c_ac_cat and movements in total 
    :param dict taffic_kwargs: kwargs for reading traffic data
    :param str labels: List with strings to identify the traffics, ignorred if table is used as input 
    :param dict weight_classes: Text per VVC weight class to aggregate the traffic
    :param str xlabel: label for the x-axis
    :param str ylabel: label for the y-axis
    :param str|function tickformat: format string or function for ticklabels
    :param set ylim: range for the y-axis
    :param str fname: Name for the file, default is '' and no fig will be saved
    :param set figsize: Figsize in inches, default (21/2.54, 7/2.54)
    :param int dpi: dpi for saving figure to file, default is 600
    :return: if fname='', return a Matplotlib figure and axes.
    """
    
    if table is None:
        table = table_aircraft_types(traffic=traffic,
                                     traffic_kwargs=traffic_kwargs,
                                     weight_classes=weight_classes,
                                     labels=labels)    
    # plot
    return plot_bar(table,
                    xlabel=xlabel,
                    ylabel=ylabel,
                    ylim=ylim,
                    tickformat=tickformat,
                    figsize=figsize,
                    dpi=dpi,
                    fname=fname,
                    wordtable=wordtable,
                    **kwargs)


class BracketPlot(object):
    def __init__(self, slond_colors=None, figsize=None, capacity_color='#cdbbce', takeoff_color='#4a8ab7',
                 landing_color='#fdbb4b'):
        # Create an ID
        self.id = str(pd.Timestamp.utcnow())

        # Set the plotting options
        self.figsize = (21 / 2.54, 21 / 2.54) if figsize is None else figsize
        self.capacity_color = capacity_color
        self.slond_colors = slond_colors if slond_colors is not None else {
            'S': '#cdbbce',
            'L': '#cdbbce',
            'O': '#8fbbd6',
            'N': '#d1e6bd',
            'D': '#f1b7b1'
        }
        self.landing_color = landing_color
        self.takeoff_color = takeoff_color

        # Create a new figure
        self.fig, self.ax = self.create_figure()

        # Create a placeholder for the capacity bracket
        self.capacity_bracket = None

    def create_figure(self):
        """
        Initialize a new bracket plot.

        :return: the figure and axis handles.
        """

        # Create a new figure with figsize
        fig, ax = plt.subplots(num=self.id, figsize=self.figsize)
        fig.set_size_inches(21 / 2.54, 9 / 2.54)

        return fig, ax

    def add_capacity_bracket(self, bracket=None):
        self.select()

        if bracket is None:
            bracket = self.capacity_bracket
        else:
            self.capacity_bracket = bracket

        # Get the x-coordinates for the brackets
        x = bracket.data.columns + .5

        # Add the SLOND colors
        if 'period' in bracket.data.index:
            color = bracket.data.loc['period', :].map(self.slond_colors).values
        else:
            color = self.capacity_color

        # Plot the takeoffs on top and the landings on the bottom
        self.ax.bar(x, -bracket.data.loc['L', :], width=0.92, color=color, edgecolor=None)
        self.ax.bar(x, bracket.data.loc['T', :], width=0.92, color=color, edgecolor=None)

    def add_traffic_bracket(self, bracket):
        self.select()

        x = bracket.data

        # Plot the takeoffs on top and the landings on the bottom
        self.ax.bar(x.columns + .5, -x.loc['L', :], width=0.5, facecolor=self.takeoff_color, edgecolor='#757575',
                    lw=0.25)
        self.ax.bar(x.columns + .5, x.loc['T', :], width=0.5, facecolor=self.landing_color, edgecolor='#757575',
                    lw=0.25)

        # Set the lines of the grid
        self.ax.set_axisbelow(True)
        self.ax.set_xticks(x.columns, minor=True)
        self.ax.xaxis.grid(which='minor')
        self.ax.yaxis.grid(which='major')

        # Set the ticks of the axes
        self.ax.axes.tick_params(axis='both', which='both', labelsize=6, labelrotation=0, labelcolor='#757575', pad=4)

        # Set the x-axis
        self.ax.set_xlim([0, 72])
        self.ax.set_xlabel('')  # size = 14)                    # verberg as-label
        self.ax.xaxis.set_tick_params(which='both', length=0)  # en geen tickmarks
        self.ax.xaxis.set_major_locator(ticker.FixedLocator(np.arange(1.5, 71, 3)))
        self.ax.xaxis.set_major_formatter(ticker.FuncFormatter(lambda u, v: '{:d}:00'.format(int(u // 3))))

        # Create a second x-axis to separate the hours
        ax2 = self.ax.twiny()
        ax2.spines["bottom"].set_position(('outward', 3))
        ax2.xaxis.set_major_locator(ticker.FixedLocator(np.arange(0, 72.5, 3) / 72))
        ax2.xaxis.set_ticks_position("bottom")
        ax2.xaxis.set_tick_params(which='major', length=7, width=0.5, color='#757575')
        ax2.xaxis.set_ticklabels([])
        for side in ax2.spines:
            ax2.spines[side].set_color('none')

        # Set the y-axis
        self.ax.yaxis.set_tick_params(which='both', length=0)
        self.ax.yaxis.set_major_locator(ticker.ticker.FormatStrFormatter(5))
        self.ax.yaxis.set_major_formatter(ticker.FuncFormatter(lambda u, v: '{:1.0f}'.format(abs(u))))

    def add_takeoff_landing_label(self, takeoff_label='takeoffs', landing_label='landings'):
        self.select()

        ylim = self.ax.get_ylim()

        ymid = -ylim[0] / (-ylim[0] + ylim[1])
        for y, label in zip([ymid / 2, (1 + ymid) / 2], [landing_label, takeoff_label]):
            self.ax.text(-0.041, y,
                         label,
                         ha='right',
                         va='center',
                         rotation='vertical',
                         transform=self.ax.transAxes)
        for y1, y2 in zip([0, ymid], [ymid, 1]):
            line = Line2D([-0.035, -0.035], [0.02 + y1, y2 - 0.02],
                          lw=0.5, color='#757575', transform=self.ax.transAxes)
            line.set_clip_on(False)
            self.ax.add_line(line)

    def add_capacity_legend(self, label='runway capacity'):
        w = 1 / 72
        x = 0.72
        ys = [1.0492, 1.0353, 1.0492, 1.0561, 1.0457]
        heights = [4 * w, 4 * w, 3 * w, 2 * w, 3.5 * w]
        for s, y, h in zip(self.slond_colors, ys, heights):
            if 'period' in self.capacity_bracket.data.index:
                c = self.slond_colors[s]
                self.ax.text(x + w / 2, 1.07, s, fontsize=4, ha='center', va='center', transform=self.ax.transAxes)
            else:
                c = self.capacity_color

            rect = Rectangle(xy=(x, y), width=w, height=h, facecolor=c, edgecolor='white', linewidth=0.5, clip_on=False,
                             transform=self.ax.transAxes)
            self.ax.add_patch(rect)
            x += w

        self.ax.text(x + w / 2, 1.07, label, fontsize=6, ha='left', va='center', transform=self.ax.transAxes)

    def add_traffic_legend(self, label='traffic'):

        self.select()

        # Add traffic legend
        w = .5 / 72
        y = 1.07
        h1, h2 = [6 * w, 2 * w, 3 * w], [-2 * w, -4 * w, -7 * w]
        colors = [self.landing_color, self.takeoff_color]
        for c, heights in zip(colors, [h1, h2]):
            x = 0.92
            for h in heights:
                rect = Rectangle(xy=(x, y), width=w, height=h, facecolor=c, edgecolor='#757575', linewidth=0.5,
                                 clip_on=False, transform=self.ax.transAxes)
                self.ax.add_patch(rect)
                x += w

        self.ax.text(x + w / 2, 1.07, label, fontsize=6, ha='left', va='center', transform=self.ax.transAxes)

    def select(self):
        plt.figure(self.id)
        plt.sca(self.ax)

    def save(self, *args, **kwargs):
        return self.fig.savefig(*args, **kwargs)

    def show(self):
        return self.fig.show()

# -----------------------------------------------------------------------------
# Get colors
# -----------------------------------------------------------------------------
def get_cycler_color(index=0):
    ''' Cycler color o.b.v. de index
    '''
    return plt.rcParams['axes.prop_cycle'].by_key()['color'][index]

# -----------------------------------------------------------------------------
# legenda
# -----------------------------------------------------------------------------
def update_legend_position(ax, target, fig=None):
    '''Update legend position, target, o.b.v. bbox.x1
    '''
    if fig is None:
        fig = plt.gcf()
    transf = fig.transFigure.inverted()
    renderer = fig.canvas.get_renderer()

    bbox = ax.get_tightbbox(renderer).transformed(transf)
    diff = target-bbox.x1

    l, b, w, h = ax.get_position().bounds
    # corrigeer x-positie
    ax.set_position([l+diff, b, w, h])


def get_text_bbox(handle, ax=None, fig=None):
    '''Bounding box of text
    '''
    if fig is None: fig = plt.gcf()
    if ax is None: ax = plt.gca()

    fig.canvas.draw()
    return handle.get_window_extent().inverse_transformed(ax.transData)

     
def plot_line(table,
              x,
              y,
              xlabel=None,
              ylabel=None,
              xtickformat=None,
              ytickformat=None,
              xstep=None,
              ystep=None,
              ncol=None, ###Todo: via xParams
              clip_on=False,
              fname='',                 
              figsize=(8.27, 2.76),
              dpi=600,
              wordtable=None,
              **kwargs):
    """
    :param pd.DataFrame table: Dataframe to plot, see table_aircraft_types   
    :param int|str x: the column name of the data to visualise
    :param int|str y: the column name of the data to visualise
    :param str xlabel: label for the x-axis
    :param str ylabel: label for the y-axis
    :param str|function xtickformat: format string or function for ticklabels
    :param str|function ytickformat: format string or function for ticklabels
    :param float|None xstep: step value for the x-axis
    :param float|None ystep: step value for the y-axis
    :param int ncol: number of columns in legend
    :param boolean clip_on: clipping on plot area, default False    
    :param str fname: Name for the file, default is '' and no fig will be saved
    :param set figsize: Figsize in inches, default (21/2.54, 7/2.54)
    :param int dpi: dpi for saving figure to file, default is 600
    :return: if fname='', return a Matplotlib figure and axes.
    """
    
    
    # Plot
    ax = table.plot(x=x,
                    y=y,
                    figsize=figsize,
                    clip_on=clip_on,
                    **kwargs)

    # X-as
    if xtickformat is not None:
        if isinstance(xtickformat, str):
            ax.xaxis.set_major_formatter(ticker.StrMethodFormatter(xtickformat))
        elif callable(xtickformat):
            ax.xaxis.set_major_formatter(ticker.FuncFormatter(xtickformat))
        else:
            print('Ignore unknown xtickformat')

    if len(table.index.names) == 2:
        ax.set_xticklabels(table.index.get_level_values(1))
        branding.set_xlabels(table.index.get_level_values(0).unique(), ax=ax)
    elif xlabel is not None:
        branding.set_xlabels(xlabel, ax=ax)
    else:
        ax.set_xlabel('') # verberg as-label

    if xstep is not None:
        ax.xaxis.set_major_locator(ticker.MultipleLocator(xstep))

    # Y-as
    if ytickformat is not None:
        if isinstance(ytickformat, str):
            ax.yaxis.set_major_formatter(ticker.StrMethodFormatter(ytickformat))
        elif callable(ytickformat):
            ax.yaxis.set_major_formatter(ticker.FuncFormatter(ytickformat))
        else:
            print('Ignore unknown tickformat')

    if ylabel is not None:
        branding.set_ylabels(ylabel, ax=ax)

    if ystep is not None:
        ax.yaxis.set_major_locator(ticker.MultipleLocator(ystep))

    # legend
    if ncol is None: ncol = len(y)
    ax.legend(ncol=ncol, **branding.xParams['legend'])

    # save figure
    fig = plt.gcf()
    if fname:
        fig.savefig(fname, dpi=dpi)
    
    # Export figure to Word
    if wordtable:
        fig_to_word(fig=fig,
                    table=wordtable,
                    dpi=dpi)

    if not (fname or wordtable):
        return fig, ax 
    else:
        plt.close(fig)
        return


def plot_history(history, 
                 history_kwargs = {'sheet_name': 'realisatie'},
                 x='jaar',
                 y='verkeer',
                 label='realisatie',
                 xlabel='jaar',
                 ylabel='vliegtuigbewegingen',
                 xtickformat='{x:.0f}',
                 ytickformat=Formatter_1000sep0d,
                 xstep=1,
                 ystep=None,
                 ncol=None, ###Todo: via xParams
                 clip_on=False,                   
                 dpi=600,
                 fname='',
                 wordtable=None,                 
                 figsize=(8.27, 2.76),
                 **kwargs):

    """
    :param pd.DataFrame|str history: the historic dataset to visualise, should contain the specified x and y as columns.
    If it is a string then the file will be inported in a pd.DataFrame.
    :param dict history_kwargs: optional arguments for read_file
    :param int|str x: the column name of the data to visualise, defaults to 'jaar'.
    :param int|str y: the column name of the data to visualise, defaults to 'verkeer'.
    :param str label: label for the legend
    :param str xlabel: label for the x-axis
    :param str ylabel: label for the y-axis
    :param float|None xstep: step value for the x-axis
    :param float|None ystep: step value for the y-axis
    :param int ncol: number of columns in legend
    :param boolean clip_on: clipping on plot area, default False    
    :param int dpi: dpi for saving figure to file, default is 600
    :param str fname: Name for the file, default is '' and no fig will be saved
    :param set figsize: Figsize in inches, default (21/2.54, 7/2.54)
    :return: if fname='', return a Matplotlib figure and axes.
    """

    # Import history data in dataframe
    if isinstance(history, str):
        history = read_file(history, **history_kwargs)

    # Plot
    return plot_line(table=history,
                     x=x,
                     y=y,
                     label=label,
                     xlabel=xlabel,
                     ylabel=ylabel,
                     xtickformat=xtickformat,
                     ytickformat=ytickformat,
                     xstep=xstep,
                     ystep=ystep,
                     ncol=ncol, ###Todo: via xParams
                     clip_on=clip_on,
                     fname=fname,                 
                     figsize=figsize,
                     wordtable=wordtable,
                     **kwargs)


def plot_prediction(history, 
                    prediction=None,
                    stats=None,
                    history_kwargs = {'sheet_name': 'realisatie'},
                    x='jaar',
                    y='verkeer',
                    labels=['realisatie', 'prognose'],
                    xlabel='jaar',
                    ylabel='vliegtuigbewegingen',
                    xtickformat='{x:.0f}',
                    ytickformat=Formatter_1000sep0d,
                    xstep=1,
                    ystep=None,
                    ncol=None, ###Todo: via xParams
                    clip_on=False,                   
                    dpi=600,
                    fname='',                 
                    figsize=(8.27, 2.76),
                    wordtable=None,
                    **kwargs):

    """
    :param pd.DataFrame|str history: the historic dataset to visualise, should contain the specified x and y as columns.
    If it is a string then the file will be inported in a pd.DataFrame.
    :param pd.DataFrame prediction: the raw predicted values, should contain the specified column_name as the data and a
    x-column. Will be superseeded if stats are given
    :param pd.DataFrame stats: the min, max and mean predicted values and year as index.
    :param dict history_kwargs: optional arguments for read_file
    :param int|str x: the column name of the data to visualise, defaults to 'jaar'.
    :param int|str y: the column name of the data to visualise, defaults to 'verkeer'.
    :param str label: label for the legend
    :param str xlabel: label for the x-axis
    :param str ylabel: label for the y-axis
    :param float|None xstep: step value for the x-axis
    :param float|None ystep: step value for the y-axis
    :param int ncol: number of columns in legend
    :param boolean clip_on: clipping on plot area, default False    
    :param int dpi: dpi for saving figure to file, default is 600
    :param str fname: Name for the file, default is '' and no fig will be saved
    :param set figsize: Figsize in inches, default (21/2.54, 7/2.54)
    :return: if fname='', return a Matplotlib figure and axes.
    """

    # Import history data in dataframe
    if isinstance(history, str):
        history = read_file(history, **history_kwargs)

    # Plot the history
    fig, ax = plot_history(history=history,
                           x=x,
                           y=y,
                           label=labels[0],
                           xlabel=xlabel,
                           ylabel=ylabel,
                           xtickformat=xtickformat,
                           ytickformat=ytickformat,
                           xstep=xstep,
                           ystep=ystep,
                           ncol=ncol, ###Todo: via xParams
                           clip_on=clip_on,
                           fname='',                 
                           figsize=figsize,
                           zorder=3, # plot on top
                           wordtable=None,
                           **kwargs)

    # Describe the prediction for each year
    if stats is None:
        stats = prediction.groupby(x)[y].describe()
    means = stats['mean']
    lo = stats['min']
    hi = stats['max']
    
    # Combine last point from history with prediction
    history = history.set_index(x)
    p_mean = history[y].tail(1).append(means)
    p_lo = history[y].tail(1).append(lo)
    p_hi = history[y].tail(1).append(hi)

    # Color the background of the prediction
    ax.fill_between(p_mean.index,
                    p_lo,
                    p_hi,
                    clip_on=clip_on,
                    **branding.xParams['prediction_fill'])
    
    # Plot the error bars
    e = ax.errorbar(means.index,
                    means,
                    yerr=[means-lo, hi-means],  # asymmetrisch interval
                    label=labels[1],
                    clip_on=clip_on, #Bug in Matplotlib? Lijkt niet te werken
                    **branding.xParams['errorbar'])
    # clip_on?
    if not clip_on:
        for b in e[1]:
            b.set_clip_on(False)
    
    # Plot de lijnen voor het gemiddelde
    ax.plot(p_mean.index,
            p_mean,
            color=get_cycler_color(1),
            marker='None',
            clip_on=clip_on)
    
    # legend
    if ncol is None: ncol = len(y)
    ax.legend(ncol=ncol, **branding.xParams['legend'])

    # save figure
    fig = plt.gcf()
    if fname:
        fig.savefig(fname, dpi=dpi)

    # Export figure to Word
    if wordtable:     
        fig_to_word(fig=fig,
                    table=wordtable,
                    dpi=dpi)

    if not (fname or wordtable):
        return fig, ax 
    else:
        plt.close(fig)
        return


def plot_prediction2(history, prediction, column_name='data', prediction_errorbar_kwargs=None,
                    prediction_fill_kwargs=None, history_plot_kwargs=None,doc29_factor=None):
    """

    :param pd.DataFrame history: the historic dataset to visualise, should contain the specified column_name as the data
    and a 'year' column.
    :param pd.DataFrame prediction: the predicted values, should contain the specified column_name as the data and a
    'year' column.
    :param int|str column_name: the column name of the data to visualise, defaults to 'data'.
    :param dict history_plot_kwargs: argument arguments to overwrite the settings used for visualising the historic data.
    :param dict prediction_errorbar_kwargs: arguments to overwrite the settings used for visualising the errorbars of
    the prediction.
    :param dict prediction_fill_kwargs: arguments to overwrite the settings used for visualising the filled area
    of the prediction.
    :return: a Matplotlib figure and axes.
    """
    # Apply the custom history plot style if provided
    history_style = {'marker': 'o', 'markeredgewidth': 2, 'fillstyle': 'none', 'label': 'history',
                     'color': '#141251'}
    if history_plot_kwargs is not None:
        history_style.update(history_plot_kwargs)

    # Apply the custom prediction errobar style if provided
#    prediction_style = {'marker': '_', 'capsize': 4, 'ecolor': '#9491AA', 'markeredgewidth': 4,
#                        'markeredgecolor': '#9491AA', 'fillstyle': 'none', 'color': '#1B60DB', 'label': 'prediction'}
    if prediction_errorbar_kwargs is not None:
        prediction_style.update(prediction_errorbar_kwargs)

    # Apply the custom prediction fill_between style if provided
    prediction_fill = {'color': '#027E9B', 'alpha': 0.3}
    if prediction_fill_kwargs is not None:
        prediction_fill.update(prediction_fill_kwargs)


    # Create a figure
    fig, ax = plt.subplots(figsize=(12, 4))

    # Plot the history
    plt.plot(history['years'], history[column_name], **history_style)
    
    # Describe the prediction for each year
    statistics = prediction[column_name]/doc29_factor

    # Plot the prediction
    plt.errorbar(history['years'].tail(1).tolist() + [prediction['years'].max()],
                 history[column_name].tail(1).tolist() + [statistics[1]],
                 yerr=[[0] + [(statistics[1]- statistics[0])],
                       [0] + [(statistics[2] - statistics[1])]],
                 **prediction_style)
    
    # Color the background of the prediction
    plt.fill_between(history['years'].tail(1).tolist() + [prediction['years'].max()],
                     history[column_name].tail(1).tolist() + [statistics[0]],
                     history[column_name].tail(1).tolist() + [statistics[2]],
                     **prediction_fill)


    # Set the xticks
    ax.set_xticks(np.arange(history['years'].min(), prediction['years'].max() + 1, 1))

    # Add horizontal grid lines
    ax.grid(axis='y')
    ax.set_ylim(bottom=0)
    # Add a legend
    plt.legend(ncol=2, bbox_to_anchor=(0.9, 1.15))
    
    if doc29_factor:
        ax.set_ylabel('NRM', color='k')  
    
    
        scale_factor=doc29_factor
        
        ax2 = ax.twinx()  # instantiate a second axes that shares the same x-axis
        
        ax2.plot(history['years'],scale_factor*history[column_name], alpha=0)
        
        ax2.set_ylabel('Doc. 29')  # we already handled the x-label with ax1
#        ax2.tick_params(axis='y', labelcolor='k')
        ax2.set_ylim(bottom=0)
    
    return fig, ax


def plot_windrose(windrose):
    # Get the directions
    directions = windrose.index.get_level_values(0).unique()
    directions = directions[directions > 0]

    # Get the directions
    speeds = windrose.index.get_level_values(1).unique()
    speeds = speeds[speeds > 0]

    # Fill zero values, necessary for bar stacking
    idx = pd.MultiIndex.from_product([directions, speeds])
    data = windrose.reindex(idx)

    # Calculate the angles
    theta = np.deg2rad(directions)

    # Calculate the percentages and reshape the data
    data_x = (data / windrose.sum() * 100).reset_index().pivot(index='level_0', columns='level_1', values='STN')
    bottom = np.zeros(data_x.shape) + 5
    bottom[:, 1:] += data_x.iloc[:, :-1].cumsum(axis=1).values
    bottom_x = pd.DataFrame(bottom, index=data_x.index, columns=data_x.columns)

    # Create a polar axis system
    ax = plt.subplot(111, projection='polar', facecolor='0.9')
    ax.set_theta_zero_location("N")
    ax.set_theta_direction(-1)
    ax.set_ylim(0, 30)
    ax.set_yticks([5, 15, 25, 35])
    ax.set_yticklabels(['', '10 %', '20 %', '30 %'])
    ax.set_xticks(theta)
    ax.set_xticklabels(['', '', 'O', '', '', 'Z', '', '', 'W', '', '', 'N'])

    for max_speed in data_x.columns:

        # Calculate the minimum speed for this maximum speed
        min_speed = max_speed - 5

        # Set the label
        if max_speed != data_x.columns.max():
            label = '{}-{} kts'.format(min_speed, max_speed)
        else:
            label = '>{} kts'.format(min_speed)

        # Remove the nan values to avoid runtime warnings
        theta_s = theta[~bottom_x[max_speed].isna()]
        data_s = data_x.loc[~bottom_x[max_speed].isna(), max_speed]
        bottom_s = bottom_x.loc[~bottom_x[max_speed].isna(), max_speed]

        # Calculate the width of the bar
        bar_width = 2 * np.pi / 12 * (0.15 + max_speed / 50)

        # Plot the bars
        p1 = ax.bar(theta_s, data_s, bottom=bottom_s, width=bar_width, label=label)

    ax.bar(0, 5, width=2 * np.pi, label='other')
    plt.text(0, 0, '{:.1f} %'.format(windrose.loc[0, 0] / windrose.sum() * 100), horizontalalignment='center',
             verticalalignment='center')
    ax.legend(loc=4)
    plt.show()


def plot_runway_usage(traffic,
                      labels=None,
                      den=('D','E','N'),
                      n=7,
                      runways=None,
                      ylabel='vliegtuigbewegingen',
                      ylim=[0,110000],
                      dy=10000,
                      reftraffic=1,
                      numbers=False,
                      fname=None,
                      wordtable=None,
                      dpi=600):
    """
    Plot runway usage per relevant runway
    
    :param str|pd.DataFrame|TrafficAggregate traffic: traffic data containing landing/type, runway, DEN and nubmer of movements
    :param str labels: List with strings to identify the scenarios. 
    :param tuple den: period of day for which to plot runway usage. 
    :param int n: number of runways to plot.
    :param str runways: list of strings to indicate order of appearance of runways.
    :param str ylabel: label for the y-axis.
    :param int ylim: list of integers to indicate range for the y-axis
    :param int dy: step size of y-axis
    :param int reftraffic: indicatie reference scenario when comparing multiple scenario's
    :param boolean numbers: if True indicatie scenario numbers, default False    
    :param str fname: Name for the file, default is '' and no fig will be saved
    :param int dpi: dpi for saving figure to file, default is 600
    :return: png-image if fname='', else return a Matplotlib figure and axes.
    """
    
    def NumberFormatter(x, pos):
        'The two args are the value and tick position'
        return '{:,.0f}'.format(x).replace(',', '.')
    def GetVal(var, i):
        'Scalar of een list'
        if isinstance(var, list):
            i = min(i, len(var)-1) # hergebruik van de laatste waarde
            return var[i]
        else:
            return var

    # kopie van plotformat
    MarkerWidth =  [0.36, 0.36, 0.2, 0.2, 0.1]  # voor 1, 2 en >2 traffics
    MarkerHeight = [0.12, 0.12, 0.1, 0.1, 0.08]
    BarWidth = [0.16, 0.16, 0.06, 0.06, 0.04]
    BarGap = [0, 0, 0.05, 0.05, 0.04]

    # converteer naar list
    if not isinstance(traffic, list): traffic = [traffic]
    if not isinstance(labels, list): labels = [labels]

    # X-positie van de bars
    x = np.arange(n)
    ntrf = len(traffic)
    i = ntrf - 1
    w = GetVal(BarWidth, i) * n/7  # normaliseer voor de aslengte
    g = GetVal(BarGap, i)          # of /ntrf?

    dx = [(w+g)*(x - 0.5*(ntrf-1)) for x in range(ntrf)]

    # markers en staafjes
    marker_height = (GetVal(MarkerHeight, i) * (ylim[-1] - ylim[0]) / 10)
    mw = GetVal(MarkerWidth, i) * n/7
    dxm = list(dx)

    # clip marker
    if ntrf == 2:
        mw = (mw + w)/2
        dxm[0] = dx[0] - (mw-w)/2
        dxm[1] = dx[1] + (mw-w)/2
    elif ntrf > 2:
        mw = min(mw, w+g)

    # twee aansluitende subplots
    fig, (ax1, ax2) = plt.subplots(1, 2, sharey=True)
    fig.set_size_inches(21/2.54, 10/2.54)

    # margins
    fig.subplots_adjust(bottom=0.18, wspace=0)

    # verwerken traffics
    realization = [] # als Casper realisaties worden gebruikt
    for i, trf in enumerate(traffic):
        
        # Read traffic
        if isinstance(trf, str):
            df = read_file(trf, sep=None, engine='python')
            if 'C_runway' in df.columns:
                trf = Traffic.read_casper_file(df)
            elif 'd_runway' in df.columns:
                trf = Traffic.read_daisy_meteoyear_file(df)

        # realisatietraffic aanpassen
        if 'C_runway' in trf.data:
            realization.append(i)
            trf.add_den()
            trf.add_landing_takeoff()
            trf.data = (trf.data.rename(columns={'LT':'d_lt','C_runway':'d_runway','DEN':'d_den'})
                                .assign(d_myear=2019, total=1))

        # Get the runway statistics
        trf_stats = trf.get_runway_usage_statistics('|'.join(den)).reset_index()

        # sorteer
        if 'key' not in trf_stats.columns:
            trf_stats['key'] = trf_stats['d_lt'] + trf_stats['d_runway']

        if runways is not None:
            # tweede traffic in dezelfde volgorde
            keys = [k + r for k in runways for r in runways[k]]    # keys: combinatie van lt en runway
            sorterIndex = dict(zip(keys, range(len(keys))))        # plak een volgnummer aan de keys
            trf_stats['order'] = trf_stats['key'].map(sorterIndex) # soteerindex toevoegen
            trf_stats = trf_stats.sort_values(by=['order'])        # sorteer dataframe
        else:
            trf_stats = trf_stats.sort_values(by=['d_lt', 'mean'], ascending=False)
            runways = {'L': trf_stats['d_runway'].loc[trf_stats['d_lt'] == 'L'],
                       'T': trf_stats['d_runway'].loc[trf_stats['d_lt'] == 'T']}

        # maak de plot
        for lt, xlabel, ax in zip(['T', 'L'], ['starts', 'landingen'], [ax1, ax2]):

            # selecteer L of T
            trf2 = trf_stats.loc[trf_stats['d_lt'] == lt]
            trf2 = trf2.head(n) # gaat alleen goed als er ook echt n-runways zijn

            # staafjes
            bar_height = trf2['max'] - trf2['min']
            if i == reftraffic:
                c = 1
            else:
                c = 0

            ax.bar(x+dx[i],
                   height=bar_height.values,
                   bottom=trf2['min'].values,
                   width=w,
                   color=get_cycler_color(c))

            # gemiddelde
            ax.bar(x+dxm[i],
                    height=marker_height,
                    bottom=trf2['mean'].values-marker_height/2,
                    width=mw,
                    color=get_cycler_color(c))

            # opmaak, alleen de eerste keer
            if i == 0:
                # geen vertikale gridlines
                ax.grid(which='major', axis='x', b=False)

                # X-as
                ax.margins(x=0.02)
                ax.set_xticks(x)
                ax.set_xticklabels(trf2['d_runway'])
                branding.set_xlabels(xlabel, gap=0.02, ax=ax)

                # Y-as
                ax.set_ylim(ylim)
                ax.yaxis.set_major_locator(ticker.MultipleLocator(base=dy))
                ax.yaxis.set_major_formatter(ticker.FuncFormatter(NumberFormatter))
                if ax==ax1:
                    branding.set_ylabels(ylabel, ax=ax)


    # scheidingslijntje tussen subplots
    for ax, side, i in [[ax1, 'right', 1], [ax2, 'left', 0]]:
        ax.spines[side].set_color('none')
        ax.axvline(ax.get_xlim()[i],
                   marker='',
                   color='white',
                   linewidth=2,
                   zorder=0)

    # nummertjes
    if numbers:
        xp = []
        yp = []
        for i, p in enumerate(ax1.patches):
            if not i%(n*2):
                xp.append(p.get_x() + p.get_width()/2)
                yp.append(p.get_y())
        for i, x in enumerate(xp):
            ax1.text(x, 0.95 * min(yp), str(i),
                     ha='center', va='top', fontsize=3)

    # legenda
    if ntrf == 1:
        w *= 0.8      # legenda op 80%
        g *= 0.8
        mw *= 0.8
        dx = [0]      # x voor de staafjes
        dxm = [0]     # x voor de markers
        plot_matrix = [(0, 0.30, 0.05, 0.2, 'left')]        
    elif ntrf == 2:
        w *= 0.8                # legenda op 80%
        g *= 0.8
        mw *= 0.8
        dx = [-w/2, w/2]        # x voor de staafjes
        dxm = [-mw/2, mw/2]     # x voor de markers
        plot_matrix = [(0, 0.3, 0.05, -0.2, 'right'),
                       (1, 0.4, 0.15, 0.2, 'left')]
    else:
        dx = [(w+g)* x for x in range(4)]
        dxm = dx
        reftraffic=3
        plot_matrix = [(0, 0.30, 0.05, 0.15, None),
                       (1, 0.25, 0.00, 0.15, None),
                       (2, 0.35, 0.10, 0.15, 'left'),
                       (3, 0.30, 0.05, 0.15, 'left')]

    # plotgebied
    ax0 = fig.add_axes([0, 0.89, 0.025, 0.1])

    # geen assen
    ax0.axis('off')

    # genormaliseerde asses
    ax0.set_xlim(0, 0.5*n/7)
    ax0.set_ylim(0, 1)

    for i, yi, bottom, xt, ha in plot_matrix:
        if i == reftraffic:
            c = 1
        else:
            c = 0
        if i not in realization:
            ax0.bar(dx[i], height=0.5, bottom=bottom,
                    width=w,
                    color=get_cycler_color(c),
                    clip_on=False)
        ax0.bar(dxm[i], height=0.05, bottom=yi,
                width=mw,
                color=get_cycler_color(c),
                clip_on=False)
        if ha:
            t = ax0.text(dx[i]+xt , 0.50, labels[c],
                         horizontalalignment=ha,
                         verticalalignment='top')
        if i == 2:
            # maak ruimte voor text label
            bbox = get_text_bbox(t)
            dx[3] += bbox.x1


    # Check uitlijning van de legend
    update_legend_position(ax0, target=0.9-0.01)

    # Save figure
    if fname:
        fig.savefig(fname, dpi=dpi)

    # Export figure to Word
    if wordtable:
        fig_to_word(fig=fig,
                    table=wordtable,
                    dpi=dpi)
        
    if not (fname or wordtable):
        return fig, (ax1, ax2) 
    else:
        plt.close(fig)
        return
    


def plot_noise_bba(grids,
                   scale=1.0,
                   levels=[48, 58],
                   noise='Lden',
                   mean='mean',
                   refine_factor=10,
                   labels=['gemiddeld', 'weersinvloeden'],
                   figsize=(21/2.54, 21/2.54),
                   fname=None,
                   dpi=600,
                   wordtable=None,
                   **kwargs):
    """
    Create a bandwidth plot showing the noise contours of various scenarios or meteo years.
    
    :param str|Grid grids: Either a folder location containing envira-files, or a MultiGrid object
    :param float scale: Scaling factor apllied to the noise grids
    :param int levels: List with integers to clarify which dB-values to plot.
    :param str noise: For Lden or Lnight grids 
    :param str mean: Average grid type to use for the average noise levels
    :param integer refine_factor: Factor to refine the grid with bi-cubic spline interpolation
    :param str labels: List labels for average noise level and bandwith in the legend
    :param set figsize: Figsize in inches, default (21/2.54, 21/2.54)
    :param str fname: (Optional) Name for the file to save. Default is None and no fig will be saved but fig, ax is returned
    :param int dpi: dpi for saving figure to file, default is 600
    :return: if fname='' saved image, else return a Matplotlib figure and axes.
    """
   
    # Get the noise grids
    if isinstance(grids, str):
        grids = Grid.read_enviras(grids, noise=noise).scale(scale)
    
    # Initialize plot
    plot = GridPlot(grids, figsize=figsize, **kwargs)

    # Make the plot   
    plot.add_bandwidth(levels=levels,
                       labels=labels,
                       mean=mean,
                       refine_factor=refine_factor)
        
    # Free memory
    gc.collect()
    
    # Save figure
    if fname:
        plot.save(fname, dpi=dpi)
        
    # Export figure to Word
    if wordtable:
        fig_to_word(fig=plot.fig,
                    table=wordtable,
                    dpi=dpi)

    if not (fname or wordtable):
        return plot.fig, plot.ax 
    else:
        plt.close(plot.fig)
        return
    
    
def plot_noise_diff(grid,
                    other_grid,
                    scale=1.0,
                    levels=[48,58],
                    colors=None,
                    deltas=[-1.5,1.5],
                    noise='Lden',
                    mean='mean',
                    refine_factor=10,
                    labels=['Scenario 1','Scenario 2'],
                    figsize=(21/2.54, 21/2.54),
                    fname=None,
                    dpi=600,
                    wordtable=None,
                    **kwargs):
    
    """
    Create a difference plot showing two noise contours including a heatmap.
    
    :param str|Grid grid: either a folder location containing envira-files, or a MultiGrid object
    :param str|Grid other_grid: either a folder location containing envira-files, or a MultiGrid object
    :param float scale: the scaling factor. Standard set to 1.0
    :param int levels: List with integers to clarify which dB-values to plot.
    :param str noise: For Lden or Lnight grids 
    :param str mean: Average grid type to use for the average noise levels
    :param integer refine_factor: Factor to refine the grid with bi-cubic spline interpolation
    :param str labels: List with strings to identify the scenarios. 
    :param str fname: (Optional) Name for the file to save. Default is None and no fig will be saved but fig, ax is returned
    :param int dpi: dpi for saving figure to file, default is 600
    :return: if fname='' saved image, else return a Matplotlib figure and axes.
    """

    # Default colors
    if colors is None:
        colors = [get_cycler_color(i) for i in range(2)]
        
    # Get the noise grids
    grid = read_grid(grid, noise=noise, mean=mean)
    other_grid = read_grid(other_grid, noise=noise, mean=mean)
                   
    # initialize plot
    plot = GridPlot(grid, 
                    other_grid=other_grid,
                    figsize=figsize,
                    **kwargs)

    # add heatmap
    plot.add_comparison_heatmap(other_grid,
                                deltas=deltas)
    # add colorbar
    plot.add_colorbar()
    
    # add contour lines
    plot.add_contours(levels=levels)    
    plot.add_contours(levels=levels, other_grid=True, contourlabels=False)
        
    # Add legend
    if labels is not None:
        cs1 = Line2D([], [], color=colors[0], marker='None')
        cs2 = Line2D([], [], color=colors[1], marker='None')
        plot.ax.legend(handles=[cs1, cs2],
                       labels=labels,
                       title=r'Geluidbelasting $L_{' + grid.unit[1:] + r'}$',
                       **branding.xParams['contourlegend'])                        

    # Free memory
    gc.collect()
    
    # Save figure
    if fname:
        plot.save(fname, dpi=dpi)
        
    # Export figure to Word
    if wordtable:
        fig_to_word(fig=plot.fig,
                    table=wordtable,
                    dpi=dpi)
        
    if not (fname or wordtable):
        return plot.fig, plot.ax 
    else:
        plt.close(plot.fig)
        return
 


def plot_iaf_sec(traffic,
                 routesector=dir+'/data/RouteSector.txt',
                 svg_template=dir+'/data/FigSectorisatie_template.svg',
                 fname='fig/figure_24.svg',
                 wordtable=None):
    
    """
    Create a plot for the sector and routes
    
    :param str|TrafficAggregate traffic: traffic data containing information about route and movements in total 
    :param str routesector: File with route names and corresponding sector or iaf 
    :param str svg_template: File with svg definition for the figure 
    :param str fname: (Optional) Name for the file to save. Default is None and no fig will be saved but fig, ax is returned
    :return: return a svg-file.
    """
    
    # if traffic is string, read file, else keep traffic as TrafficAggregate
    if isinstance(traffic, str):
        traffic = Traffic.read_daisy_mean_file(traffic)
    
    # add sector to traffic  
    rs = pd.read_csv(routesector,sep='\t')
    traffic.add_sector(rs)
    
    # get distribution
    sector = traffic.get_sector_distribution()
        
    # Normalise the results
    sids = rs[rs['lt']=='T']['sector'].unique().tolist()
    stars = rs[rs['lt']=='L']['sector'].unique().tolist()
    sector[sids] = 100 * sector[sids] / sector[sids].sum(axis=0)
    sector[stars] = 100 * sector[stars] / sector[stars].sum(axis=0)
    sector = round(sector,0)
    
    data=pd.DataFrame(sector)
        
    data = data.rename(columns={'total':'Value'})
    data['Length']= data['Value']*7
    
    # read template
    with open(svg_template, 'r') as svgt:
        xmlcontents = svgt.read()
    
    # fill in values and lengths
    for iaf_sec in sids + stars:
        xmlcontents = xmlcontents.replace("p"+iaf_sec, str(data.at[iaf_sec,'Value'])).replace("R"+iaf_sec, str(data.at[iaf_sec,'Length']))

    # colors landingen and starts
    xmlcontents = xmlcontents.replace("color_1", get_cycler_color(3))
    xmlcontents = xmlcontents.replace("color_2", get_cycler_color(1))
    
    # Write svg-file
    svg_file = Path(fname).with_suffix('.svg')
    with open(svg_file, 'w') as svg:
        svg.write(xmlcontents)
    
    # convert svg
    if Path(fname).suffix != '.svg':
        os.system(f'inkscape -C --export-width=2480 "{svg_file}" --export-png="{fname}"')

    # Export figure to Word
    if wordtable:
        fig_to_word(fig=fname,
                    table=wordtable)
        
def plot_noise_heatmap(grid,
                   scale=1.0,
                   levels=[48, 58],
                   refine_factor=10,
                   figsize=(21/2.54, 21/2.54),
                   fname=None,
                   dpi=600,
                   wordtable=None,
                   **kwargs):
    
    # Create a figure
    plot = GridPlot(grid, 
                    figsize=figsize,
                    **kwargs,
                    # background=background[i],
                    # extent=size[i]['extent'], 
                    # xlim = size[i]['xlim'], 
                    # ylim = size[i]['ylim'],
                       )
    
    # Add the heatmap
    plot.add_heatmap(vmin=levels[0], vmax=levels[1],colormap=plt.cm.get_cmap('RdYlGn_r'),alpha=0.3,refine_factor=3) #Spectral_r
    
    # Add a colorbar
    plot.add_colorbar(cax_position=[0.85, 0.8, 0.05, 0.17])
    
    # Free memory
    gc.collect()
    
    # Save figure
    if fname:
        plot.save(fname, dpi=dpi)
        
    # Export figure to Word
    if wordtable:
        fig_to_word(fig=plot.fig,
                    table=wordtable,
                    dpi=dpi)

    if not (fname or wordtable):
        return plot.fig, plot.ax 
    else:
        plt.close(plot.fig)
        return